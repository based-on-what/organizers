#!/usr/bin/env python3
"""
Enhanced Page Counter for Documents

This script counts pages in PDF, EPUB, and DOCX files with improved error handling
and modern library support. MOBI support has been removed due to deprecated dependencies.
"""

import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional

try:
    from pypdf import PdfReader  # Modern replacement for PyPDF2
except ImportError:
    try:
        from PyPDF2 import PdfReader  # Fallback to old library
        print("Warning: Using deprecated PyPDF2. Please install pypdf instead.")
    except ImportError:
        print("Error: Neither pypdf nor PyPDF2 is available. Install with: pip install pypdf")
        sys.exit(1)

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    print("Error: ebooklib is required. Install with: pip install ebooklib")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

from shared_utils import (
    setup_logging, safe_file_operation, save_results_to_file, 
    find_files_by_extensions, ProgressReporter
)

def count_pdf_pages(file_path: Path) -> int:
    """Count pages in a PDF file using modern pypdf library."""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            return len(reader.pages)
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")


def count_epub_pages(file_path: Path) -> int:
    """Count approximate pages in an EPUB file."""
    try:
        book = epub.read_epub(str(file_path))
        # Count HTML documents as pages
        return len(list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT)))
    except Exception as e:
        raise Exception(f"Error processing EPUB: {str(e)}")


def count_docx_pages(file_path: Path) -> int:
    """
    Count pages in a DOCX file.
    Note: This is an approximation based on page breaks and content.
    """
    try:
        doc = Document(str(file_path))
        
        # Method 1: Count explicit page breaks
        page_breaks = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '\f' in run.text:  # Form feed character indicates page break
                    page_breaks += 1
                    
        # Method 2: Estimate based on content if no explicit breaks found
        if page_breaks == 0:
            total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            # Rough estimation: 2000 characters per page
            estimated_pages = max(1, total_chars // 2000)
            return estimated_pages
            
        return page_breaks + 1  # Add 1 for the first page
        
    except Exception as e:
        raise Exception(f"Error processing DOCX: {str(e)}")


def count_file_pages(file_path: Path) -> int:
    """Count pages in a file based on its extension."""
    extension = file_path.suffix.lower()
    
    page_counters = {
        '.pdf': count_pdf_pages,
        '.epub': count_epub_pages,
        '.docx': count_docx_pages,
    }
    
    counter = page_counters.get(extension)
    if not counter:
        raise Exception(f"Unsupported file format: {extension}")
        
    return counter(file_path)

def count_pages_in_directory(directory: Path) -> List[Tuple[str, int]]:
    """
    Count pages in supported files in the specified directory.
    
    Args:
        directory: Directory to scan for supported files
        
    Returns:
        List of tuples containing (filename, page_count)
    """
    # Supported file extensions (removed .mobi and .doc due to deprecated libraries)
    supported_extensions = {'.pdf', '.epub', '.docx'}
    
    # Find all supported files
    files = find_files_by_extensions(directory, supported_extensions, recursive=False)
    
    if not files:
        logging.info("No supported files found in the directory.")
        return []
    
    files_with_pages = []
    error_files = []
    
    # Set up progress reporting
    progress = ProgressReporter(len(files), "Processing files")
    logging.info(f"Found {len(files)} supported files to process")
    
    for file_path in files:
        try:
            # Use shared utility for safe file operations
            if not safe_file_operation(file_path, "page counting"):
                error_files.append((file_path.name, "File access error"))
                progress.update()
                continue
            
            # Count pages
            page_count = count_file_pages(file_path)
            files_with_pages.append((file_path.name, page_count))
            logging.info(f"✓ {file_path.name}: {page_count} pages")
            
        except Exception as e:
            error_msg = str(e)
            logging.error(f"✗ Error processing {file_path.name}: {error_msg}")
            error_files.append((file_path.name, error_msg))
        
        progress.update()
    
    progress.finish()
    
    # Report errors if any
    if error_files:
        logging.warning(f"\nSummary of {len(error_files)} files with errors:")
        for filename, error in error_files:
            logging.warning(f"  {filename}: {error}")
    
    return files_with_pages

def display_and_save_results(files_with_pages: List[Tuple[str, int]], output_file: str = "page_count_results.txt") -> None:
    """
    Display results and save them to a file.
    
    Args:
        files_with_pages: List of (filename, page_count) tuples
        output_file: Output file name
    """
    if not files_with_pages:
        logging.info("No files were successfully processed.")
        return
    
    # Sort files by page count
    sorted_files = sorted(files_with_pages, key=lambda x: x[1])
    
    # Display results
    logging.info("\nFiles sorted by page count:")
    total_pages = 0
    for filename, pages in sorted_files:
        logging.info(f"{filename}: {pages} pages")
        total_pages += pages
    
    logging.info(f"\nSummary:")
    logging.info(f"Total files: {len(sorted_files)}")
    logging.info(f"Total pages: {total_pages}")
    
    # Save results using shared utility
    results_text = []
    for filename, pages in sorted_files:
        results_text.append(f"{filename}: {pages} pages")
    
    results_text.extend([
        "",
        f"Total files: {len(sorted_files)}",
        f"Total pages: {total_pages}"
    ])
    
    if save_results_to_file(results_text, Path(output_file), "FILES SORTED BY PAGE COUNT"):
        logging.info(f"Results saved to '{output_file}'")


def main() -> None:
    """Main function to execute the page counting process."""
    # Set up logging
    logger = setup_logging("INFO")
    
    # Get current directory
    current_directory = Path.cwd()
    logger.info(f"Scanning directory: {current_directory}")
    
    # Count pages in files
    files_with_pages = count_pages_in_directory(current_directory)
    
    # Display and save results
    display_and_save_results(files_with_pages)


if __name__ == "__main__":
    main()
