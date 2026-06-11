"""Tests for readers.pages — dispatch logic plus tiny generated fixtures."""
import zipfile

import pytest

from readers.pages import (
    _is_image,
    count_cbz_pages,
    count_docx_pages,
    count_epub_pages,
    count_pages,
    count_pdf_pages,
    supported_extensions,
)


class TestIsImage:
    @pytest.mark.parametrize("name", [
        "page1.jpg", "page1.JPEG", "cover.png", "a.gif", "b.bmp", "c.webp",
        "dir/inside/page.PNG",
    ])
    def test_images(self, name):
        assert _is_image(name)

    @pytest.mark.parametrize("name", [
        "readme.txt", "ComicInfo.xml", "noextension", "archive.zip", "page.jpg.bak",
    ])
    def test_non_images(self, name):
        assert not _is_image(name)


class TestDispatch:
    def test_supported_extensions(self):
        assert supported_extensions() == frozenset(
            {'.pdf', '.epub', '.cbz', '.cbr', '.docx'}
        )

    def test_unsupported_raises(self, tmp_path):
        target = tmp_path / "file.xyz"
        target.write_text("data")
        with pytest.raises(RuntimeError, match="Unsupported format"):
            count_pages(target)

    def test_dispatch_case_insensitive(self, tmp_path):
        # .CBZ must route to the CBZ counter
        target = tmp_path / "comic.CBZ"
        with zipfile.ZipFile(target, 'w') as zf:
            zf.writestr("p1.jpg", b"x")
        assert count_pages(target) == 1


@pytest.fixture
def two_page_pdf(tmp_path):
    from pypdf import PdfWriter
    target = tmp_path / "doc.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.add_blank_page(width=72, height=72)
    with open(target, 'wb') as f:
        writer.write(f)
    return target


class TestPdf:
    def test_two_pages(self, two_page_pdf):
        assert count_pdf_pages(two_page_pdf) == 2

    def test_corrupt_pdf_raises_runtime_error(self, tmp_path):
        bad = tmp_path / "bad.pdf"
        bad.write_bytes(b"not a pdf at all")
        with pytest.raises(RuntimeError, match="Error reading PDF"):
            count_pdf_pages(bad)


class TestCbz:
    def test_counts_only_images(self, tmp_path):
        target = tmp_path / "comic.cbz"
        with zipfile.ZipFile(target, 'w') as zf:
            zf.writestr("p1.jpg", b"x")
            zf.writestr("p2.png", b"x")
            zf.writestr("sub/p3.webp", b"x")
            zf.writestr("ComicInfo.xml", b"<x/>")
        assert count_cbz_pages(target) == 3

    def test_corrupt_raises(self, tmp_path):
        bad = tmp_path / "bad.cbz"
        bad.write_bytes(b"not a zip")
        with pytest.raises(RuntimeError, match="Error reading CBZ"):
            count_cbz_pages(bad)


_CONTAINER_XML = """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>"""

_OPF_XML = """<?xml version="1.0"?>
<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="id">
  <metadata/>
  <manifest>
    <item id="c1" href="ch1.xhtml" media-type="application/xhtml+xml"/>
    <item id="c2" href="ch2.xhtml" media-type="application/xhtml+xml"/>
    <item id="css" href="style.css" media-type="text/css"/>
  </manifest>
  <spine><itemref idref="c1"/><itemref idref="c2"/></spine>
</package>"""


@pytest.fixture
def minimal_epub(tmp_path):
    target = tmp_path / "book.epub"
    with zipfile.ZipFile(target, 'w') as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr("META-INF/container.xml", _CONTAINER_XML)
        zf.writestr("content.opf", _OPF_XML)
        zf.writestr("ch1.xhtml", "<html/>")
        zf.writestr("ch2.xhtml", "<html/>")
        zf.writestr("style.css", "")
    return target


class TestEpub:
    def test_counts_xhtml_manifest_items(self, minimal_epub):
        assert count_epub_pages(minimal_epub) == 2

    def test_corrupt_raises(self, tmp_path):
        bad = tmp_path / "bad.epub"
        bad.write_bytes(b"not a zip")
        with pytest.raises(RuntimeError, match="Error reading EPUB"):
            count_epub_pages(bad)


class TestDocx:
    def test_short_doc_is_one_page(self, tmp_path):
        from docx import Document
        target = tmp_path / "doc.docx"
        d = Document()
        d.add_paragraph("hello world")
        d.save(str(target))
        assert count_docx_pages(target) == 1

    def test_explicit_page_breaks(self, tmp_path):
        from docx import Document
        from docx.enum.text import WD_BREAK
        target = tmp_path / "doc.docx"
        d = Document()
        p = d.add_paragraph()
        run = p.add_run("page one")
        run.add_break(WD_BREAK.PAGE)
        p.add_run("page two")
        d.save(str(target))
        assert count_docx_pages(target) == 2
